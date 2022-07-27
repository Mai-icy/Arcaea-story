#!/usr/bin/env python
# -*- coding:utf-8 -*-
from docx import Document
from pathlib import Path
import json
import re

VN_MAIN_PATH = "vn/vn-main"
VN_SIDE_PATH = "vn/vn-side"
VNS_PATH = "vns"


def parse_vns(chapter):
    chapter_path = Path(VNS_PATH) / chapter

    chapter_dict = {}
    for i in chapter_path.iterdir():
        language = re.split(r"_|\.", i.name)[-2]
        lines = i.read_text(encoding="utf-8").split("\n\n")
        res_list = []
        for line in lines:
            if not line.startswith("say "):
                continue
            line = line.strip()
            res_list.append(line[5:-1].replace("\\", ""))

        text = "\n\n".join(res_list)
        chapter_dict[language] = text
    return chapter_dict


def add_chapter(document: Document, chapter_index_tuple, file_json, languages=("en", "zh-Hans")):
    for chapter in chapter_index_tuple:
        if chapter in file_json:
            # 普通剧情在vn文件内
            story_json = file_json[chapter]
        else:
            # CG剧情文件在vns文件内
            story_json = parse_vns(chapter)

        document.add_heading(chapter, level=3)
        for lang in languages:
            document.add_heading(lang, level=4)
            document.add_paragraph(story_json[lang])


def main(languages=None):
    main_file_path = Path(VN_MAIN_PATH)
    side_file_path = Path(VN_SIDE_PATH)

    main_story_json = json.loads(main_file_path.read_text(encoding="utf-8"))
    side_story_json = json.loads(side_file_path.read_text(encoding="utf-8"))

    document = Document()
    document.add_heading('Arcaea Story', 0)
    document.add_heading('Main Story', level=1)

    document.add_heading("Hikari's Story", level=2)
    chapter_index = ("1-" + str(i) for i in range(1, 10))
    add_chapter(document, chapter_index, main_story_json, languages)

    document.add_heading("Tairitsu's Story", level=2)
    chapter_index = ("2-" + str(i) for i in range(1, 10))
    add_chapter(document, chapter_index, main_story_json, languages)

    document.add_heading("The meeting of Hikari and Tairitsu", level=2)
    chapter_index = ("100-" + str(i) for i in range(1, 6))
    add_chapter(document, chapter_index, main_story_json, languages)

    document.add_heading("Hikari versus Tairitsu", level=2)
    chapter_index = ("101-" + str(i) for i in range(1, 9))
    add_chapter(document, chapter_index, main_story_json, languages)

    document.add_heading("The End", level=2)
    chapter_index = ("102-" + str(i) for i in range(1, 8))
    add_chapter(document, chapter_index, main_story_json, languages)

    document.add_heading("The End", level=2)
    chapter_index = ("102-" + str(i) for i in range(1, 8))
    add_chapter(document, chapter_index, main_story_json, languages)

    document.add_heading("The ending", level=2)
    chapter_index = ("103-" + str(i) for i in range(1, 3))
    add_chapter(document, chapter_index, main_story_json, languages)

    document.add_heading('Side Story', 1)

    document.add_heading("Saya's Story", level=2)
    chapter_index = ("3-" + str(i) for i in range(1, 7))
    add_chapter(document, chapter_index, side_story_json, languages)

    document.add_heading("Kou's Story", level=2)
    chapter_index = ("4-" + str(i) for i in range(1, 9))
    add_chapter(document, chapter_index, side_story_json, languages)

    document.add_heading("Lethe's Story", level=2)
    chapter_index = ("5-" + str(i) for i in range(1, 7))
    add_chapter(document, chapter_index, side_story_json, languages)

    document.add_heading("Alice&Tenniel's Story", level=2)
    chapter_index = ("7-" + str(i) for i in range(1, 7))
    add_chapter(document, chapter_index, side_story_json, languages)

    document.add_heading("Lagrange's Story", level=2)
    chapter_index = ("9-" + str(i) for i in range(1, 7))
    add_chapter(document, chapter_index, side_story_json, languages)

    document.add_heading("Lagrange's Story", level=2)
    chapter_index = ("99-" + str(i) for i in range(1, 7))
    add_chapter(document, chapter_index, side_story_json, languages)

    document.add_heading("Eto&luna's Story", level=2)
    chapter_index = ("10-" + str(i) for i in range(1, 7))
    add_chapter(document, chapter_index, side_story_json, languages)

    document.add_heading("Shirabe's Story", level=2)
    chapter_index = ("6-" + str(i) for i in range(1, 4))
    add_chapter(document, chapter_index, side_story_json, languages)

    document.add_heading("Mir's Story", level=2)
    chapter_index = ("6-" + str(i) for i in range(1, 4))
    add_chapter(document, chapter_index, side_story_json, languages)

    document.add_heading("Ayu's Story", level=2)
    chapter_index = ("11-" + str(i) for i in range(1, 4))
    add_chapter(document, chapter_index, side_story_json, languages)

    document.add_heading("Vita's Story", level=2)
    chapter_index = ("12-" + str(i) for i in range(1, 4))
    add_chapter(document, chapter_index, side_story_json, languages)

    document.save('Arcaea-story.docx')


main()
